"""
Spec-doc ingest test suite.

Cycles 1-5  test DocxRuleParser in isolation (no network).
Cycles 6-9  test the graph write path against a real Neo4j instance.
"""
import pytest
from pathlib import Path


# ══════════════════════════════════════════════════════════════════════════════
# PARSER TESTS (no Neo4j)
# ══════════════════════════════════════════════════════════════════════════════

class TestDocxRuleParser:

    # ── Cycle 1 — tracer bullet ───────────────────────────────────────────────

    def test_br_rows_extracted_with_normalised_ids(self, rule_path, sample_docx):
        """BR04, BR 23 → BR04, BR23; BR07 stays BR07."""
        from src.docx_generic_parser import DocxRuleParser
        items, _ = DocxRuleParser(rule_path).parse(sample_docx)
        ids = {i.req_id for i in items}
        assert ids == {"BR04", "BR23", "BR07"}

    # ── Cycle 2 ───────────────────────────────────────────────────────────────

    def test_title_taken_from_first_body_line(self, rule_path, sample_docx):
        from src.docx_generic_parser import DocxRuleParser
        items, _ = DocxRuleParser(rule_path).parse(sample_docx)
        br04 = next(i for i in items if i.req_id == "BR04")
        assert br04.title == "First line of BR04"

    # ── Cycle 3 ───────────────────────────────────────────────────────────────

    def test_category_signals_matched_per_body(self, rule_path, sample_docx):
        from src.docx_generic_parser import DocxRuleParser
        items, _ = DocxRuleParser(rule_path).parse(sample_docx)
        by_id = {i.req_id: i for i in items}

        assert "SQLView"   in by_id["BR04"].candidate_categories
        assert "OracleEBS" in by_id["BR23"].candidate_categories
        assert "BatchJob"  in by_id["BR07"].candidate_categories
        assert "Java"      in by_id["BR07"].candidate_categories

    # ── Cycle 4 ───────────────────────────────────────────────────────────────

    def test_named_extractions_deduped_and_sorted(self, rule_path, sample_docx):
        from src.docx_generic_parser import DocxRuleParser
        items, _ = DocxRuleParser(rule_path).parse(sample_docx)
        br04 = next(i for i in items if i.req_id == "BR04")

        assert br04.named_extractions["views"]  == ["VW_INVOICE_HDR"]
        assert br04.named_extractions["fields"] == ["INVOICE_DATE"]

    # ── Cycle 5 ───────────────────────────────────────────────────────────────

    def test_context_collects_non_br_content_only(self, rule_path, sample_docx):
        from src.docx_generic_parser import DocxRuleParser
        _, context = DocxRuleParser(rule_path).parse(sample_docx)

        assert "introduction paragraph" in context
        assert "eInvoice"               in context   # from info table
        # BR table rows must NOT appear in context
        assert "BR04"  not in context
        assert "BR 23" not in context

    # ── Cycle 10 ──────────────────────────────────────────────────────────────

    def test_multiple_tables_with_different_allowlisted_prefixes(self, rule_path, multi_prefix_docx):
        """A document with a BRU table and a separate BRM table extracts both, IDs preserved verbatim."""
        from src.docx_generic_parser import DocxRuleParser
        items, _ = DocxRuleParser(rule_path).parse(multi_prefix_docx)
        ids = {i.req_id for i in items}
        assert ids == {"BRU01", "BRU23", "BRM01", "BRM23"}

    # ── Cycle 11 ──────────────────────────────────────────────────────────────

    def test_unrecognized_prefix_warns_instead_of_silently_dropping(self, rule_path, tmp_path):
        """An ID shaped like a BR row but with a prefix not in the allow-list is skipped and warned about, not extracted."""
        from docx import Document as DocxDocument
        from src.docx_generic_parser import DocxRuleParser

        doc = DocxDocument()
        table = doc.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "BRX05"
        table.cell(0, 1).text = "This uses a prefix that was never added to the rule file."
        path = tmp_path / "unrecognized_prefix.docx"
        doc.save(str(path))

        parser = DocxRuleParser(rule_path)
        items, _ = parser.parse(path)

        assert items == []
        assert any("BRX05" in w for w in parser.warnings)

    # ── Cycle 12 ──────────────────────────────────────────────────────────────

    def test_plain_prefix_not_shadowed_by_longer_prefix(self, rule_path, tmp_path):
        """A plain 'BR' table alongside a 'BRU' table: BR04 must not be swallowed by the BRU alternative."""
        from docx import Document as DocxDocument
        from src.docx_generic_parser import DocxRuleParser

        doc = DocxDocument()

        br_table = doc.add_table(rows=1, cols=2)
        br_table.cell(0, 0).text = "BR04"
        br_table.cell(0, 1).text = "Plain BR rule body."

        bru_table = doc.add_table(rows=1, cols=2)
        bru_table.cell(0, 0).text = "BRU01"
        bru_table.cell(0, 1).text = "Utility rule body."

        path = tmp_path / "prefix_of_prefix.docx"
        doc.save(str(path))

        items, _ = DocxRuleParser(rule_path).parse(path)
        ids = {i.req_id for i in items}
        assert ids == {"BR04", "BRU01"}

    # ── Cycle 13 ──────────────────────────────────────────────────────────────

    def test_brw_prefix_is_allowlisted(self, rule_path, tmp_path):
        """Regression: a BRM table alongside a BRW table extracts both — BRW must be in id_prefixes."""
        from docx import Document as DocxDocument
        from src.docx_generic_parser import DocxRuleParser

        doc = DocxDocument()

        brm = doc.add_table(rows=2, cols=2)
        brm.cell(0, 0).text = "BRM01"
        brm.cell(0, 1).text = "BRM rule one."
        brm.cell(1, 0).text = "BRM23"
        brm.cell(1, 1).text = "BRM rule two."

        brw = doc.add_table(rows=2, cols=2)
        brw.cell(0, 0).text = "BRW01"
        brw.cell(0, 1).text = "BRW rule one."
        brw.cell(1, 0).text = "BRW23"
        brw.cell(1, 1).text = "BRW rule two."

        path = tmp_path / "brm_brw.docx"
        doc.save(str(path))

        parser = DocxRuleParser(rule_path)
        items, _ = parser.parse(path)

        ids = {i.req_id for i in items}
        assert ids == {"BRM01", "BRM23", "BRW01", "BRW23"}
        assert parser.warnings == []


# ══════════════════════════════════════════════════════════════════════════════
# GRAPH TESTS  (require Neo4j at bolt://localhost:7687)
# ══════════════════════════════════════════════════════════════════════════════

class TestDocumentHierarchy:

    # shared helpers ──────────────────────────────────────────────────────────

    def _ingest(self, graph_db, items, context):
        """Mirror the endpoint's code path: ParsedItem → BR/Document → graph."""
        from models import BR as _BR, Document as _Doc, UseCase as _UC

        doc_model = _Doc(
            uc_id="UC99", doc_type="FDD", flow_name="TestFlow",
            source_file="test_fixture.docx", context=context,
        )
        br_models = [
            _BR(
                br_id=item.req_id, uc_id="UC99", doc_type="FDD",
                flow_name="TestFlow", title=item.title, body=item.body,
                candidate_categories=item.candidate_categories,
                affected_views=item.named_extractions.get("views", []),
                affected_fields=item.named_extractions.get("fields", []),
                source_file=item.source_file,
            )
            for item in items
        ]
        graph_db.upsert_flows({"TestFlow"})
        graph_db.upsert_use_cases([_UC(uc_id="UC99", project_id="", flow_name="TestFlow")])
        graph_db.upsert_documents([doc_model])
        graph_db.upsert_brs(br_models)
        return doc_model.node_id

    # ── Cycle 6 — tracer bullet ───────────────────────────────────────────────

    def test_document_node_created_with_context(self, graph_db, rule_path, sample_docx):
        """Document node has context, uc_id, doc_type, source_file set."""
        from src.docx_generic_parser import DocxRuleParser
        items, context = DocxRuleParser(rule_path).parse(sample_docx, source_label="UC99")

        doc_id = self._ingest(graph_db, items, context)

        with graph_db.driver.session() as s:
            row = s.run("MATCH (d:Document {id: $id}) RETURN d", id=doc_id).single()

        assert row is not None
        d = row["d"]
        assert d["doc_type"]    == "FDD"
        assert d["uc_id"]       == "UC99"
        assert d["source_file"] == "test_fixture.docx"
        assert "introduction"   in d["context"]

    # ── Cycle 7 ───────────────────────────────────────────────────────────────

    def test_br_nodes_linked_to_document_via_defines(self, graph_db, rule_path, sample_docx):
        """3 BR nodes are reachable from Document via [:DEFINES]."""
        from src.docx_generic_parser import DocxRuleParser
        items, context = DocxRuleParser(rule_path).parse(sample_docx, source_label="UC99")

        doc_id = self._ingest(graph_db, items, context)

        with graph_db.driver.session() as s:
            row = s.run(
                "MATCH (d:Document {id: $id})-[:DEFINES]->(r:BR) RETURN count(r) AS cnt",
                id=doc_id,
            ).single()

        assert row["cnt"] == 3

    # ── Cycle 8 ───────────────────────────────────────────────────────────────

    def test_usecase_node_linked_to_document(self, graph_db, rule_path, sample_docx):
        """UseCase→[:HAS_DOCUMENT]→Document edge is created."""
        from src.docx_generic_parser import DocxRuleParser
        items, context = DocxRuleParser(rule_path).parse(sample_docx, source_label="UC99")

        doc_id = self._ingest(graph_db, items, context)

        with graph_db.driver.session() as s:
            row = s.run(
                """
                MATCH (uc:UseCase {uc_id: 'UC99'})-[:HAS_DOCUMENT]->(d:Document {id: $id})
                RETURN uc
                """,
                id=doc_id,
            ).single()

        assert row is not None
        assert row["uc"]["flow_name"] == "TestFlow"

    # ── Cycle 9 ───────────────────────────────────────────────────────────────

    def test_flow_node_linked_to_usecase(self, graph_db, rule_path, sample_docx):
        """Flow→[:HAS_UC]→UseCase edge is created (matches CLI pipeline schema)."""
        from src.docx_generic_parser import DocxRuleParser
        items, context = DocxRuleParser(rule_path).parse(sample_docx, source_label="UC99")

        self._ingest(graph_db, items, context)

        with graph_db.driver.session() as s:
            row = s.run(
                """
                MATCH (f:Flow {name: 'TestFlow'})-[:HAS_UC]->(uc:UseCase {uc_id: 'UC99'})
                RETURN f
                """
            ).single()

        assert row is not None
