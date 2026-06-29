"""
Shared fixtures for the webapp test suite.

Parser tests:  use sample_docx + rule_path — no network required.
Graph tests:   use graph_db — requires Neo4j at bolt://localhost:7687.
"""
import sys
from pathlib import Path

import pytest
from docx import Document as DocxDocument

# ── Path setup ────────────────────────────────────────────────────────────────
# Add webapp/ so `from src.X import ...` works, and pipeline/src/ so graph.py
# can import models and config without modification.
_REPO   = Path(__file__).parent.parent.parent
_WEBAPP = _REPO / "webapp"
_PIPE   = _REPO / "pipeline" / "src"

for p in [str(_WEBAPP), str(_PIPE)]:
    if p not in sys.path:
        sys.path.insert(0, p)

RULE_PATH  = _REPO / "parsing-rules" / "br_requirements.yml"
NEO4J_URI  = "neo4j://127.0.0.1:7687"
NEO4J_USER = "neo4j"
NEO4J_PASS = "nguyenngochai"


# ── Parser fixtures ───────────────────────────────────────────────────────────

@pytest.fixture
def rule_path() -> Path:
    return RULE_PATH


@pytest.fixture
def sample_docx(tmp_path) -> Path:
    """
    Minimal .docx with known content:
      - one intro paragraph  (→ context)
      - one non-BR info table (→ context)
      - one BR table with three rows
    """
    doc = DocxDocument()

    # Context: paragraph
    doc.add_paragraph("This is the introduction paragraph.")

    # Context: non-BR info table
    info = doc.add_table(rows=2, cols=2)
    info.cell(0, 0).text = "Project"
    info.cell(0, 1).text = "eInvoice"
    info.cell(1, 0).text = "Version"
    info.cell(1, 1).text = "1.0"

    # BR table
    brs = doc.add_table(rows=3, cols=2)
    brs.cell(0, 0).text = "BR04"
    brs.cell(0, 1).text = (
        "First line of BR04\n"
        "System reads from VW_INVOICE_HDR view.\n"
        '"INVOICE_DATE" field is mandatory.'
    )
    brs.cell(1, 0).text = "BR 23"
    brs.cell(1, 1).text = "System calls Oracle EBS to retrieve payment data."
    brs.cell(2, 0).text = "BR07"
    brs.cell(2, 1).text = "Java service triggers a batch job every night at 02:00."

    path = tmp_path / "test_fixture.docx"
    doc.save(str(path))
    return path


# ── Graph fixture ─────────────────────────────────────────────────────────────

@pytest.fixture
def graph_db():
    """Real Neo4j connection.  Cleans up test nodes after each test."""
    from graph import GraphBuilder

    db = GraphBuilder(NEO4J_URI, NEO4J_USER, NEO4J_PASS)
    yield db

    with db.driver.session() as s:
        s.run("MATCH (n) WHERE n.source_file = 'test_fixture.docx' DETACH DELETE n")
        s.run("MATCH (n:UseCase {uc_id: 'UC99'}) DETACH DELETE n")
        s.run("MATCH (n:Flow {name: 'TestFlow'}) DETACH DELETE n")

    db.close()
