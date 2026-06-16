"""
Azure DevOps Server (on-premises) MCP Server

Transport: SSE/HTTP — run on a shared team VM, one instance for all developers.

Auth: Basic (username:password base64). If your ADO Server requires NTLM,
install `requests-ntlm` and swap the auth section (see comments below).
"""

import html
import io
import json
import logging
import os
import re
from contextlib import asynccontextmanager
from typing import Any

import anyio
import httpx
import uvicorn
from dotenv import load_dotenv
from httpx_ntlm import HttpNtlmAuth
from mcp.server import Server
from mcp.server.sse import SseServerTransport
from mcp.types import TextContent, Tool
from starlette.applications import Starlette
from starlette.requests import Request
from starlette.routing import Mount, Route

load_dotenv()

logging.basicConfig(level=logging.INFO)

logger = logging.getLogger("ado-mcp")

# --------------------------------------------------------------------------
# Configuration (all from .env)
# --------------------------------------------------------------------------

ADO_BASE_URL = os.getenv("ADO_BASE_URL", "").rstrip("/")
ADO_PROJECT = os.getenv("ADO_PROJECT", "")
ADO_USERNAME = os.getenv("ADO_USERNAME", "")
ADO_PASSWORD = os.getenv("ADO_PASSWORD", "")
SERVER_HOST = os.getenv("MCP_HOST", "0.0.0.0")
SERVER_PORT = int(os.getenv("MCP_PORT", "8001"))
API_VERSION = os.getenv("ADO_API_VERSION", "5.0")

_required = [ADO_BASE_URL, ADO_PROJECT, ADO_USERNAME, ADO_PASSWORD]

if not all(_required):
    raise EnvironmentError(
        "Missing required env vars: ADO_BASE_URL, ADO_PROJECT, ADO_USERNAME, ADO_PASSWORD"
    )

# --------------------------------------------------------------------------
# HTTP client helpers
# --------------------------------------------------------------------------


def _ntlm_auth() -> HttpNtlmAuth:
    """NTLM auth for Windows domain accounts (enskcom\\username)."""
    return HttpNtlmAuth(ADO_USERNAME, ADO_PASSWORD)


def _headers() -> dict[str, str]:
    return {
        "Content-Type": "application/json",
        "Accept": "application/json",
    }


def _url(path: str) -> str:
    """Build a full API URL under the configured project."""
    return f"{ADO_BASE_URL}/{ADO_PROJECT}/_apis/{path}"


def _wi_url(path: str) -> str:
    return _url(f"wit/{path}")


# --------------------------------------------------------------------------
# ADO API calls (sync-wrapped for simplicity; run in thread via anyio)
# --------------------------------------------------------------------------


def _get(url: str, params: dict | None = None) -> dict:
    p = {"api-version": API_VERSION, **(params or {})}
    resp = httpx.get(url, auth=_ntlm_auth(), headers=_headers(), params=p, timeout=15, verify=False)
    resp.raise_for_status()
    return resp.json()


def _post(url: str, body: dict, params: dict | None = None) -> dict:
    p = {"api-version": API_VERSION, **(params or {})}
    resp = httpx.post(url, auth=_ntlm_auth(), headers=_headers(), params=p, json=body, timeout=15, verify=False)
    resp.raise_for_status()
    return resp.json()


# --------------------------------------------------------------------------
# Business logic
# --------------------------------------------------------------------------


def fetch_work_item(item_id: int) -> dict[str, Any]:
    data = _get(_wi_url(f"workitems/{item_id}"), params={"$expand": "all"})
    fields = data.get("fields", {})
    relations = data.get("relations", [])

    # Flatten the most useful fields
    result: dict[str, Any] = {
        "id": data.get("id"),
        "url": data.get("_links", {}).get("html", {}).get("href", ""),
        "type": fields.get("System.WorkItemType"),
        "title": fields.get("System.Title"),
        "state": fields.get("System.State"),
        "assigned_to": _display_name(fields.get("System.AssignedTo")),
        "created_by": _display_name(fields.get("System.CreatedBy")),
        "created_date": fields.get("System.CreatedDate"),
        "changed_date": fields.get("System.ChangedDate"),
        "iteration": fields.get("System.IterationPath"),
        "area": fields.get("System.AreaPath"),
        "priority": fields.get("Microsoft.VSTS.Common.Priority"),
        "description": _strip_html(fields.get("System.Description", "")),
        "acceptance_criteria": _strip_html(
            fields.get("Microsoft.VSTS.Common.AcceptanceCriteria", "")
        ),
        "tags": fields.get("System.Tags", ""),
        "relations": [
            {
                "rel": r.get("rel"),
                "url": r.get("url"),
                "attributes": r.get("attributes", {}),
            }
            for r in relations
        ],
    }
    return result


def search_work_items(wiql_where: str, top: int = 20) -> list[dict[str, Any]]:
    """
    Run a WIQL query. Pass a WHERE clause fragment, e.g.:
      "[System.State] = 'Active' AND [System.IterationPath] UNDER 'Sprint 12'"
    """
    query = (
        f"SELECT [System.Id], [System.Title], [System.State], "
        f"[System.WorkItemType], [System.AssignedTo] "
        f"FROM WorkItems "
        f"WHERE [System.TeamProject] = '{ADO_PROJECT}' AND {wiql_where} "
        f"ORDER BY [System.ChangedDate] DESC"
    )
    data = _post(_wi_url("wiql"), body={"query": query}, params={"$top": top})
    items = data.get("workItems", [])
    if not items:
        return []

    ids = ",".join(str(i["id"]) for i in items[:top])
    batch = _get(
        _wi_url("workitems"),
        params={
            "ids": ids,
            "fields": "System.Id,System.Title,System.State,"
                      "System.WorkItemType,System.AssignedTo,System.IterationPath",
        },
    )
    return [
        {
            "id": w["id"],
            "type": w["fields"].get("System.WorkItemType"),
            "title": w["fields"].get("System.Title"),
            "state": w["fields"].get("System.State"),
            "assigned_to": _display_name(w["fields"].get("System.AssignedTo")),
            "iteration": w["fields"].get("System.IterationPath"),
        }
        for w in batch.get("value", [])
    ]


def fetch_comments(item_id: int) -> list[dict[str, Any]]:
    """
    Reads discussion history via the work-item revisions endpoint (api-version 5.0),
    which works on all on-prem ADO Server versions.  Each revision that has a
    non-empty System.History field is one comment / history entry.
    """
    data = _get(_wi_url(f"workitems/{item_id}/revisions"))
    results = []
    for rev in data.get("value", []):
        fields = rev.get("fields", {})
        history = _strip_html(fields.get("System.History", ""))
        if not history:
            continue
        results.append({
            "revision": rev.get("rev"),
            "author": _display_name(fields.get("System.ChangedBy")),
            "date": fields.get("System.ChangedDate"),
            "text": history,
        })
    return results


def fetch_related_items(item_id: int) -> list[dict[str, Any]]:
    """Return work-item relations with resolved titles where possible."""
    item = fetch_work_item(item_id)
    results = []
    for rel in item.get("relations", []):
        rel_type = rel.get("rel", "")
        url = rel.get("url", "")

        # Resolve only work-item links (not attachments, hyperlinks, etc.)
        if "_apis/wit/workitems/" in url:
            try:
                related_id = int(url.rstrip("/").split("/")[-1])
                related = _get(_wi_url(f"workitems/{related_id}"))
                f = related.get("fields", {})
                results.append(
                    {
                        "relation_type": rel_type,
                        "id": related_id,
                        "type": f.get("System.WorkItemType"),
                        "title": f.get("System.Title"),
                        "state": f.get("System.State"),
                    }
                )
            except Exception as e:
                logger.warning("Could not resolve related item from %s: %s", url, e)
        else:
            results.append(
                {
                    "relation_type": rel_type,
                    "url": url,
                    "attributes": rel.get("attributes", {}),
                }
            )
    return results


# --------------------------------------------------------------------------
# Utility helpers
# --------------------------------------------------------------------------


def _display_name(field: Any) -> str | None:
    if isinstance(field, dict):
        return field.get("displayName") or field.get("name")
    return field


def _strip_html(raw_html: str) -> str:
    """Naïve HTML stripper — good enough for ADO description fields."""
    if not raw_html:
        return ""
    text = re.sub(r"<[^>]+>", " ", raw_html)
    text = html.unescape(text)
    return re.sub(r"\s+", " ", text).strip()


def _read_docx(data: bytes) -> str:

    from docx import Document

    doc = Document(io.BytesIO(data))
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip(" |"):
                parts.append(row_text)
    return "\n".join(parts)


def _read_xlsx(data: bytes) -> str:

    import openpyxl

    wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    parts: list[str] = []
    for sheet in wb.worksheets:
        parts.append(f"=== Sheet: {sheet.title} ===")
        for row in sheet.iter_rows(values_only=True):
            row_str = " | ".join(str(v) if v is not None else "" for v in row)
            if row_str.strip(" |"):
                parts.append(row_str)
    wb.close()
    return "\n".join(parts)


# --------------------------------------------------------------------------
# Attachment helpers
# --------------------------------------------------------------------------


def _attachment_relations(item_id: int) -> list[dict[str, Any]]:
    """Return raw AttachedFile relations for a work item."""
    data = _get(_wi_url(f"workitems/{item_id}"), params={"$expand": "relations"})
    return [
        r for r in data.get("relations", [])
        if r.get("rel") == "AttachedFile"
    ]


def list_attachments(item_id: int) -> list[dict[str, Any]]:
    results = []
    for rel in _attachment_relations(item_id):
        attrs = rel.get("attributes", {})
        results.append({
            "name": attrs.get("name", ""),
            "size_bytes": attrs.get("resourceSize"),
            "comment": attrs.get("comment", ""),
            "url": rel.get("url", ""),
        })
    return results


def read_attachment(item_id: int, filename: str) -> dict[str, Any]:
    """
    Download the first attachment whose name matches *filename* (case-insensitive)
    and return its extracted text content.  Supports .docx and .xlsx only.
    """
    relations = _attachment_relations(item_id)
    match = next(
        (r for r in relations
         if r.get("attributes", {}).get("name", "").lower() == filename.lower()),
        None,
    )
    if match is None:
        return {"error": f"No attachment named '{filename}' found on work item {item_id}."}

    name: str = match["attributes"]["name"]
    ext = name.rsplit(".", 1)[-1].lower() if "." in name else ""
    if ext not in ("docx", "xlsx"):
        return {"error": f"Unsupported file type '.{ext}'. Only .docx and .xlsx are supported."}

    url = match["url"]
    resp = httpx.get(
        url,
        auth=_ntlm_auth(),
        params={"api-version": API_VERSION},
        timeout=60,
        verify=False,
    )
    resp.raise_for_status()
    content = _read_docx(resp.content) if ext == "docx" else _read_xlsx(resp.content)
    return {"name": name, "content": content}


# --------------------------------------------------------------------------
# MCP Server definition
# --------------------------------------------------------------------------

mcp_server = Server("ado-mcp")


@mcp_server.list_tools()
async def list_tools() -> list[Tool]:
    return [
        Tool(
            name="get_work_item",
            description=(
                "Fetch full details for an Azure DevOps work item by ID. "
                "Returns title, type, state, description, acceptance criteria, "
                "tags, priority, assignee, dates, and all relations."
            ),
            inputSchema={
                "type": "object",
                "properties": {
                    "item_id": {
                        "type": "integer",
                        "description": "The numeric work item ID (ticket number).",
                    }
                },
                "required": ["item_id"],
            },
        ),
        Tool(
            name="search_work_items",
            description=(
                "Search work items using a WIQL WHERE clause. "
                "Example clauses: \"[System.State] = 'Active'\", "
                "\"[System.WorkItemType] = 'Bug' AND [System.IterationPath] UNDER 'Sprint 5'\". "
                "Returns up to `top` results (default 20)."
            ),
            inputSchema={
                "type": "object",
                "properties": {
                    "where_clause": {
                        "type": "string",
                        "description": "WIQL WHERE fragment (do not include the WHERE keyword).",
                    },
                    "top": {
                        "type": "integer",
                        "description": "Maximum results to return (default 20, max 50).",
                        "default": 20,
                    },
                },
                "required": ["where_clause"],
            },
        ),
        Tool(
            name="get_work_item_comments",
            description=(
                "Fetch the discussion thread (comments) for a work item. "
                "Useful for recovering the reasoning behind decisions made during a ticket."
            ),
            inputSchema={
                "type": "object",
                "properties": {
                    "item_id": {
                        "type": "integer",
                        "description": "The numeric work item ID.",
                    }
                },
                "required": ["item_id"],
            },
        ),
        Tool(
            name="get_related_items",
            description=(
                "Fetch all items linked to a work item (parent, child, related, duplicate, etc.). "
                "Resolves work-item links to include title and state."
            ),
            inputSchema={
                "type": "object",
                "properties": {
                    "item_id": {
                        "type": "integer",
                        "description": "The numeric work item ID.",
                    }
                },
                "required": ["item_id"],
            },
        ),
        Tool(
            name="list_work_item_attachments",
            description=(
                "List all file attachments on a work item. "
                "Returns the filename, size, and optional comment for each attachment. "
                "Use this before read_work_item_attachment to discover what files are available."
            ),
            inputSchema={
                "type": "object",
                "properties": {
                    "item_id": {
                        "type": "integer",
                        "description": "The numeric work item ID.",
                    }
                },
                "required": ["item_id"],
            },
        ),
        Tool(
            name="read_work_item_attachment",
            description=(
                "Download and extract the text content of a .docx or .xlsx attachment "
                "from a work item. Use list_work_item_attachments first to get the exact filename. "
                "Returns the full text of the document (paragraphs and tables for .docx; "
                "all sheets/cells for .xlsx)."
            ),
            inputSchema={
                "type": "object",
                "properties": {
                    "item_id": {
                        "type": "integer",
                        "description": "The numeric work item ID.",
                    },
                    "filename": {
                        "type": "string",
                        "description": "Exact filename of the attachment (e.g. 'requirements.docx').",
                    },
                },
                "required": ["item_id", "filename"],
            },
        ),
    ]


@mcp_server.call_tool()
async def call_tool(name: str, arguments: dict) -> list[TextContent]:
    try:
        if name == "get_work_item":
            result = await anyio.to_thread.run_sync(
                lambda: fetch_work_item(int(arguments["item_id"]))
            )
        elif name == "search_work_items":
            top = min(int(arguments.get("top", 20)), 50)
            result = await anyio.to_thread.run_sync(
                lambda: search_work_items(arguments["where_clause"], top)
            )
        elif name == "get_work_item_comments":
            result = await anyio.to_thread.run_sync(
                lambda: fetch_comments(int(arguments["item_id"]))
            )
        elif name == "get_related_items":
            result = await anyio.to_thread.run_sync(
                lambda: fetch_related_items(int(arguments["item_id"]))
            )
        elif name == "list_work_item_attachments":
            result = await anyio.to_thread.run_sync(
                lambda: list_attachments(int(arguments["item_id"]))
            )
        elif name == "read_work_item_attachment":
            item_id = int(arguments["item_id"])
            filename = arguments["filename"]
            result = await anyio.to_thread.run_sync(
                lambda: read_attachment(item_id, filename)
            )
        else:
            result = {"error": f"Unknown tool: {name}"}

        return [TextContent(type="text", text=json.dumps(result, indent=2, default=str))]

    except httpx.HTTPStatusError as e:
        msg = f"ADO API error {e.response.status_code}: {e.response.text}"
        logger.error(msg)
        return [TextContent(type="text", text=json.dumps({"error": msg}))]
    except Exception as e:
        logger.exception("Tool %s failed", name)
        return [TextContent(type="text", text=json.dumps({"error": str(e)}))]


# --------------------------------------------------------------------------
# SSE / Starlette wiring
# --------------------------------------------------------------------------

sse_transport = SseServerTransport("/messages/")


async def handle_sse(request: Request):
    async with sse_transport.connect_sse(
        request.scope, request.receive, request._send
    ) as (read_stream, write_stream):
        await mcp_server.run(
            read_stream,
            write_stream,
            mcp_server.create_initialization_options(),
        )


@asynccontextmanager
async def lifespan(_app):
    logger.info(
        "ADO MCP Server running on http://%s:%s/sse  (project: %s)",
        SERVER_HOST,
        SERVER_PORT,
        ADO_PROJECT,
    )
    yield


app = Starlette(
    lifespan=lifespan,
    routes=[
        Route("/sse", endpoint=handle_sse),
        Mount("/messages/", app=sse_transport.handle_post_message),
    ],
)

if __name__ == "__main__":
    uvicorn.run(app, host=SERVER_HOST, port=SERVER_PORT)
